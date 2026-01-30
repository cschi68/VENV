import math
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pump_report():
    print("--- 泵浦計算書自動產出系統 (LibreOffice 相容) ---")
    
    # 1. 輸入基本資訊
    pump_name = input("請輸入泵浦編號 (例如: P-101A/B): ")
    flow_cmh = float(input("請輸入設計流量 (CMH, m3/hr): "))
    pipe_dia = int(input("請輸入管徑 (ψ/mm): "))
    pipe_len = float(input("請輸入直管長度 (m): "))
    friction_rate = float(input("請輸入摩擦損失率 (m/100m, 預設 7.9): ") or 7.9)
    v_head = float(input("請輸入垂直揚程 (m): "))
    extra_loss = float(input("請輸入多段過濾/外部單元壓損 (m): ") or 0)
    daily_water = float(input("請輸入每日處理水量 (CMD, m3/day): ") or 0)

    # 2. 閥件數量
    ball_v = int(input("球閥數量: ") or 0)
    check_v = int(input("逆止閥數量: ") or 0)
    elbow_90 = int(input("90°彎頭數量: ") or 0)
    elbow_45 = int(input("45°彎頭數量: ") or 0)
    y_filt = int(input("Y型過濾器數量: ") or 0)

    # 3. 邏輯計算：自動判定不同管徑的等價管長係數
    # 依文件：管徑 100ψ 時球閥取 0.2, 逆止閥取 8.7, 90°彎頭取 1.7, 45°彎頭取 0.7
    # 依文件：管徑 50ψ 以下時球閥取 0.2, 逆止閥取 4.4, 90°彎頭取 0.7 或 0.9, 45°彎頭取 0.3
    if pipe_dia >= 100:
        c_ball = 0.2; c_check = 8.7; c_elbow90 = 1.7; c_elbow45 = 0.7
    else:
        c_ball = 0.2; c_check = 4.4; c_elbow90 = 0.7; c_elbow45 = 0.3
    
    eq_valves = (ball_v * c_ball) + (check_v * c_check)
    eq_elbows = (elbow_90 * c_elbow90) + (elbow_45 * c_elbow45)
    eq_others = (y_filt * 0.3)
    
    total_eq_len = pipe_len + eq_valves + eq_elbows + eq_others
    f_loss = round(total_eq_len * (friction_rate / 100), 2)
    total_h_raw = f_loss + v_head + extra_loss
    
    # 總揚程取整 (參考文件取法：45m, 25m, 15m)
    final_head = math.ceil(total_h_raw)
    if final_head < 20: final_head = 15
    elif final_head < 30: final_head = 25
    else: final_head = math.ceil(final_head / 5) * 5

    # 4. 馬力計算
    q_min = flow_cmh / 60
    # K:傳動效率(1.1), E:泵浦效率(0.7)
    kw_calc = (0.163 * q_min * final_head * 1.1) / 0.7
    kw_final = round(kw_calc, 2)
    hp_final = round(kw_final * 1.341, 2)
    
    # 運算時間
    runtime = round(daily_water / flow_cmh, 2) if flow_cmh > 0 else 0

    # 5. 產出 Docx
    doc = Document()
    safe_name = pump_name.replace("/", "_").replace("\\", "_")
    
    doc.add_heading('泵浦計算書', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 段落一：泵浦基本規格
    doc.add_paragraph(f"泵浦名稱：{pump_name}")
    p1 = doc.add_paragraph()
    run1 = p1.add_run(f"水量 = {flow_cmh} CMH  管徑 = {pipe_dia}ψ  摩擦損失為 {friction_rate} m/100m")
    run1.bold = True
    
    # 段落二：揚程計算
    doc.add_paragraph(f"一、 揚程計算：")
    doc.add_paragraph(f"等價管長 = {total_eq_len:.2f} m")
    doc.add_paragraph(f"摩擦損失 = {total_eq_len:.2f} m * {friction_rate}/100 = {f_loss} m")
    doc.add_paragraph(f"垂直揚程 = {v_head} m")
    if extra_loss > 0:
        doc.add_paragraph(f"外部單元壓損 = {extra_loss} m")
    doc.add_paragraph(f"總揚程計算：{f_loss} + {v_head} + {extra_loss} = {total_h_raw:.2f} m，實取 {final_head} m")

    # 段落三：馬力計算
    doc.add_paragraph(f"二、 泵浦馬力：")
    doc.add_paragraph(f"P = 0.163 * Q * H * K / E  ， K：傳動效率(1.1)，E：泵浦效率(0.7)")
    doc.add_paragraph(f"馬力(KW) = 0.163 * {q_min:.3f} * {final_head} * 1.1 / 0.7 = {kw_final} KW")
    doc.add_paragraph(f"約等於 {hp_final} HP")

    # 段落四：運算時間
    doc.add_paragraph(f"三、 泵浦運轉時間計算：")
    doc.add_paragraph(f"每日處理水量 = {daily_water} CMD")
    doc.add_paragraph(f"泵浦運轉時間 = {daily_water} / {flow_cmh} = {runtime} hr")

    # 存檔
    filename = f"{safe_name}_計算書.docx"
    doc.save(filename)
    print(f"\n[成功] 檔案已儲存：{filename}")

if __name__ == "__main__":
    create_pump_report()